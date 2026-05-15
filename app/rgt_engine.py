import os
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_background(cell, hex_color):
    """Helper function to set the background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_rgt(touchpoint_data: dict) -> BytesIO:
    """Generates the Requirement Gathering Template for the Bank Team."""
    
    # Extract known variables
    wud_id = touchpoint_data.get('id', 'TBD')
    api_name = touchpoint_data.get('name', 'Unnamed Integration')
    idr_details = touchpoint_data.get('idr_details', 'Standard IDR Implementation')
    
    # Initialize Document
    doc = Document()

    # ==========================================
    # GLOBAL TYPOGRAPHY SETTINGS
    # ==========================================
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Poppins'
    style_normal.font.size = Pt(11)
    if style_normal.font.element.rPr is not None and style_normal.font.element.rPr.rFonts is not None:
        style_normal.font.element.rPr.rFonts.set(qn('w:asciiTheme'), '')
        style_normal.font.element.rPr.rFonts.set(qn('w:hAnsiTheme'), '')

    # ==========================================
    # HEADER & INSTRUCTIONS
    # ==========================================
    title = doc.add_heading('Integration Requirement Template', level=1)
    title.runs[0].font.name = 'Poppins'
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.color.rgb = RGBColor(65, 105, 225) # Corporate Blue
    
    # Critical Tracking ID (Used by the inbound parser later)
    tracker = doc.add_paragraph(f"System Reference: [WUD-ID:{wud_id}]")
    tracker.runs[0].font.bold = True
    tracker.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph("\n")
    
    # Instructions for the Bank
    instruction_heading = doc.add_paragraph("ACTION REQUIRED:")
    instruction_heading.runs[0].font.bold = True
    instruction_heading.runs[0].font.color.rgb = RGBColor(225, 29, 72) # Hot Pink
    
    doc.add_paragraph(
        "Please provide the technical specifications for the integration below. "
        "Do not alter the structure of this table. Fill in your requirements directly "
        "inside the white cells marked '[Click to type...]', save the document, and reply to the email."
    )
    doc.add_paragraph("\n")

        # Extract known variables
    wud_id = touchpoint_data.get('id', 'TBD')
    api_name = touchpoint_data.get('name', 'Unnamed Integration')
    source_system = touchpoint_data.get('source', '') or ''
    business_purpose = touchpoint_data.get('business_purpose', '') or touchpoint_data.get('business_flow', '') or ''
    td = touchpoint_data.get('techDetails', {}) or {}

    def pf(val, placeholder):
        """Pre-fill if value exists (gray/locked), else show placeholder (blue/editable)."""
        v = str(val).strip() if val else ''
        return (v, False) if v else (placeholder, True)

    schema = [
        ("Touchpoint ID",               str(wud_id),                                          False),
        ("API Name",                    td.get('apiName', '') or api_name,                    False),
        ("Business Purpose",            td.get('businessPurpose', '') or business_purpose or "[To be provided]", not bool(td.get('businessPurpose', '') or business_purpose)),
        ("Interface Type",              "API",                                                False),
        ("API Type",                    *pf(td.get('apiType', ''),           "[REST / SOAP]")),
        ("Endpoint URL / Method Name",  *pf(td.get('endpointUrl', ''),       "[/v1/resource/...]")),
        ("API Method",                  *pf(td.get('apiMethod', ''),         "[GET / POST / PUT / DELETE]")),
        ("UAT URL",                     *pf(td.get('uatUrl', ''),            "[https://uat.vendor.com/...]")),
        ("Prod URL",                    *pf(td.get('prodUrl', ''),           "[https://api.vendor.com/...]")),
        ("IP Whitelisting Required",    *pf(td.get('ipWhitelist', ''),      "[Yes / No — provide IPs]")),
        ("VPN / SSL Required",          *pf(td.get('vpnRequired', ''),      "[Yes / No]")),
        ("Authentication Type",         *pf(td.get('apiAuth', ''),           "[OAuth / Basic / API Key / JWT]")),
        ("Token / Auth URL",            *pf(td.get('authDetails', ''),       "[https://auth.vendor.com/token]")),
        ("Mandatory Headers",           *pf(td.get('mandatoryHeaders', ''), "[Authorization, Content-Type, ...]")),
        ("Certificate / mTLS Notes",    *pf(td.get('certNotes', ''),        "[Client cert details if applicable]")),
        ("Sample Request Payload",      *pf(td.get('apiReq', ''),           "[Paste JSON/XML request here...]")),
        ("Sample Response Payload",     *pf(td.get('apiRes', ''),           "[Paste JSON/XML response here...]")),
        ("Error Response Sample",       *pf(td.get('errorSample', ''),      "[Paste error response here...]")),
        ("Timeout Value",               *pf(td.get('timeout', ''),          "[e.g., 30 seconds]")),
        ("Rate Limit / TPS",            *pf(td.get('rateLimitTps', ''),     "[e.g., 100 TPS]")),
        ("Retry Mechanism",             *pf(td.get('retryMechanism', ''),   "[e.g., 3x exponential backoff]")),
        ("Correlation / Reference ID",  *pf(td.get('correlationId', ''),    "[Field name for tracking]")),
        ("Callback / Webhook Required", *pf(td.get('callbackRequired', ''), "[Yes / No — URL if Yes]")),
        ("Swagger / WSDL / Postman",    *pf(td.get('swaggerUrl', ''),       "[URL or attach with reply]")),
    ]

    rgt_table = doc.add_table(rows=1, cols=2)
    rgt_table.style = 'Table Grid'
    rgt_table.autofit = False

    # Create Headers
    hdr_cells = rgt_table.rows[0].cells
    hdr_cells[0].width = Inches(2.5)
    hdr_cells[1].width = Inches(4.5)
    
    set_cell_background(hdr_cells[0], "D9D9D9") # Gray Header
    set_cell_background(hdr_cells[1], "D9D9D9") # Gray Header
    
    hdr_left = hdr_cells[0].paragraphs[0].add_run("Requirement / Key")
    hdr_left.font.bold = True
    hdr_right = hdr_cells[1].paragraphs[0].add_run("Bank Technical Input")
    hdr_right.font.bold = True

    # Populate Schema
    for key, val, is_editable in schema:
        row = rgt_table.add_row()
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(4.5)
        
        # Left Column: The Key (Always shaded gray to indicate "Locked")
        set_cell_background(row.cells[0], "F4F4F4")
        key_run = row.cells[0].paragraphs[0].add_run(key)
        key_run.font.bold = True
        
        # Right Column: The Input area
        val_run = row.cells[1].paragraphs[0].add_run(val)
        
        if not is_editable:
            # If it's our pre-filled data, shade it gray and make text standard
            set_cell_background(row.cells[1], "F4F4F4")
            val_run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            # If it's for the bank, leave background white and make text blue so they know to overwrite it
            val_run.font.color.rgb = RGBColor(65, 105, 225)
            val_run.font.italic = True

    # ==========================================
    # FINALIZATION
    # ==========================================
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    
    return word_buffer