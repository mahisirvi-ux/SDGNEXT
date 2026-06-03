import re
from io import BytesIO
from docx import Document


# Full RGT schema — must match rgt_engine.py field keys exactly
RGT_SCHEMA = [
    {"key": "apiName",          "label": "API Name"},
    {"key": "businessPurpose",  "label": "Business Purpose"},
    {"key": "apiType",          "label": "API Type"},
    {"key": "endpointUrl",      "label": "Endpoint URL / Method Name"},
    {"key": "apiMethod",        "label": "API Method"},
    {"key": "uatUrl",           "label": "UAT URL"},
    {"key": "prodUrl",          "label": "Prod URL"},
    {"key": "ipWhitelist",      "label": "IP Whitelisting Required"},
    {"key": "vpnRequired",      "label": "VPN / SSL Required"},
    {"key": "apiAuth",          "label": "Authentication Type"},
    {"key": "authDetails",      "label": "Token / Auth URL"},
    {"key": "mandatoryHeaders", "label": "Mandatory Headers"},
    {"key": "certNotes",        "label": "Certificate / mTLS Notes"},
    {"key": "apiReq",           "label": "Sample Request Payload"},
    {"key": "apiRes",           "label": "Sample Response Payload"},
    {"key": "errorSample",      "label": "Error Response Sample"},
    {"key": "timeout",          "label": "Timeout Value"},
    {"key": "rateLimitTps",     "label": "Rate Limit / TPS"},
    {"key": "retryMechanism",   "label": "Retry Mechanism"},
    {"key": "correlationId",    "label": "Correlation / Reference ID"},
    {"key": "callbackRequired", "label": "Callback / Webhook Required"},
    {"key": "swaggerUrl",       "label": "Swagger / WSDL / Postman"},
]


def extract_bank_specifications(docx_bytes: bytes) -> dict:
    """
    Parses the returned RGT Word document and extracts technical details.
    Maps the RGT table labels to the techDetails JSON keys used by details.js.
    """
    doc = Document(BytesIO(docx_bytes))
    extracted_data = {}

    if not doc.tables:
        raise ValueError("Returned document is missing the requirement table.")

    rgt_table = doc.tables[0]

    # Maps RGT label → techDetails JSON key(s)
    # Must match EXACTLY the labels in rgt_engine.py schema
    field_map = {
        "API Name":                     ["apiName"],
        "Business Purpose":             ["businessPurpose"],
        "API Type":                     ["apiType"],
        "Endpoint URL / Method Name":   ["endpointUrl"],
        "API Method":                   ["apiMethod"],
        "UAT URL":                      ["uatUrl"],
        "Prod URL":                     ["prodUrl"],
        "IP Whitelisting Required":     ["ipWhitelist"],
        "VPN / SSL Required":           ["vpnRequired"],
        "Authentication Type":          ["apiAuth"],
        "Token / Auth URL":             ["authDetails"],
        "Mandatory Headers":            ["mandatoryHeaders"],
        "Certificate / mTLS Notes":     ["certNotes"],
        "Sample Request Payload":       ["apiReq"],
        "Sample Response Payload":      ["apiRes"],
        "Error Response Sample":        ["errorSample"],
        "Timeout Value":                ["timeout"],
        "Rate Limit / TPS":             ["rateLimitTps"],
        "Retry Mechanism":              ["retryMechanism"],
        "Correlation / Reference ID":   ["correlationId"],
        "Callback / Webhook Required":  ["callbackRequired"],
        "Swagger / WSDL / Postman":     ["swaggerUrl"],
    }

    for row in rgt_table.rows:
        if len(row.cells) < 2:
            continue

        label = row.cells[0].text.strip()
        raw_value = row.cells[1].text.strip()

        if label in field_map:
            # Remove placeholder text inside [brackets]
            clean_value = re.sub(r'\[.*?\]', '', raw_value).strip()

            if clean_value:
                for db_key in field_map[label]:
                    extracted_data[db_key] = clean_value

    # Debug log
    print("\n" + "=" * 50)
    print("PARSER EXTRACTED:")
    for key, val in extracted_data.items():
        display = f"{val[:60]}..." if len(val) > 60 else val
        print(f"   - {key}: {display}")
    print("=" * 50 + "\n")

    return extracted_data


def validate_rgt_structure(docx_bytes: bytes) -> bool:
    """
    Checks if a .docx file matches our RGT template structure.
    Returns True if it's a valid RGT document, False if it's a wrong/unrelated file.

    Validation: checks if the first table has at least 50% of our known labels
    in the left column.
    """
    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception:
        return False

    if not doc.tables:
        return False

    rgt_table = doc.tables[0]

    # Our expected labels from the RGT schema
    expected_labels = {
        "Touchpoint ID", "API Name", "Business Purpose", "API Type",
        "Endpoint URL / Method Name", "API Method", "UAT URL", "Prod URL",
        "Authentication Type", "Sample Request Payload", "Sample Response Payload"
    }

    found_labels = set()
    for row in rgt_table.rows:
        if len(row.cells) >= 2:
            label = row.cells[0].text.strip()
            if label in expected_labels:
                found_labels.add(label)

    # If at least 50% of key labels are present, it's our RGT
    match_ratio = len(found_labels) / len(expected_labels)
    is_valid = match_ratio >= 0.5

    if not is_valid:
        print(f"[RGT Validate] Document failed structure check. "
              f"Found {len(found_labels)}/{len(expected_labels)} expected labels.")

    return is_valid


def compare_rgt_fields(extracted_specs: dict) -> dict:
    """
    Compares extracted bank specifications against the full RGT schema.
    Returns a breakdown of filled vs missing fields.

    Args:
        extracted_specs: dict returned by extract_bank_specifications()

    Returns:
        {
            "filled": [{"key": ..., "label": ..., "value": ...}, ...],
            "missing": [{"key": ..., "label": ...}, ...],
            "total_fields": int,
            "filled_count": int,
            "missing_count": int,
            "completion_pct": int
        }
    """
    filled = []
    missing = []

    for field in RGT_SCHEMA:
        key = field["key"]
        value = (extracted_specs.get(key) or "").strip()

        # Check if value is empty or still a placeholder [something]
        is_placeholder = bool(re.match(r'^\[.*\]$', value)) if value else False

        if value and not is_placeholder:
            filled.append({
                "key":   key,
                "label": field["label"],
                "value": value
            })
        else:
            missing.append({
                "key":   key,
                "label": field["label"]
            })

    total = len(RGT_SCHEMA)
    filled_count = len(filled)
    missing_count = len(missing)
    completion_pct = round((filled_count / total) * 100) if total > 0 else 0

    print(f"[RGT Compare] {filled_count}/{total} filled ({completion_pct}%) | {missing_count} missing")

    return {
        "filled":         filled,
        "missing":        missing,
        "total_fields":   total,
        "filled_count":   filled_count,
        "missing_count":  missing_count,
        "completion_pct": completion_pct
    }