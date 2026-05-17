import os
import json
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
# ---> LIVE API IMPORT <---
from app.core.ai_agent import generate_wud_content
from docx.enum.table import WD_TABLE_ALIGNMENT
def set_floating_background(picture):
    """Converts a standard image into a full-page background image behind text."""
    inline = picture._inline
    
    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('relativeHeight', '0')
    anchor.set('behindDoc', '1') # Pushes image behind text
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')

    simplePos = OxmlElement('wp:simplePos')
    simplePos.set('x', '0')
    simplePos.set('y', '0')
    anchor.append(simplePos)

    positionH = OxmlElement('wp:positionH')
    positionH.set('relativeFrom', 'page')
    posOffsetH = OxmlElement('wp:posOffset')
    posOffsetH.text = '0'
    positionH.append(posOffsetH)
    anchor.append(positionH)

    positionV = OxmlElement('wp:positionV')
    positionV.set('relativeFrom', 'page')
    posOffsetV = OxmlElement('wp:posOffset')
    posOffsetV.text = '0'
    positionV.append(posOffsetV)
    anchor.append(positionV)

    # UPDATED: Force image width to 8.5 inches to eliminate the right-side white gap
    extent = OxmlElement('wp:extent')
    extent.set('cx', str(int(8.5 * 914400))) 
    extent.set('cy', str(int(11.7 * 914400)))
    anchor.append(extent)
    
    wrapNone = OxmlElement('wp:wrapNone')
    anchor.append(wrapNone)

    docPr = inline.find(qn('wp:docPr'))
    if docPr is not None:
        anchor.append(docPr)
        
    graphic = inline.find(qn('a:graphic'))
    if graphic is not None:
        anchor.append(graphic)

    parent = inline.getparent()
    parent.insert(parent.index(inline), anchor)
    parent.remove(inline)
    
def add_custom_heading(doc, text, level=1, text_r=89, text_g=89, text_b=89, line_hex="C45911"):
    """Adds a heading with a specific text color and a different bottom border line color."""
    heading = doc.add_heading(text, level=level)
    
    # 1. Force the font, size, and color for the heading text (Default: Grey #595959)
    for run in heading.runs:
        run.font.name = 'Poppins'
        run.font.color.rgb = RGBColor(text_r, text_g, text_b)
        
    # 2. XML Hack to add a bottom border line (Default: Orange #C45911)
    pPr = heading._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')  # Solid line
    bottom.set(qn('w:sz'), '12')       # 1.5 pt thickness
    bottom.set(qn('w:space'), '4')     # Padding between text and line
    bottom.set(qn('w:color'), line_hex) # Use the separate line color
    
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    return heading

def set_cell_background(cell, hex_color):
    """Helper function to set the background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)
def insert_dynamic_toc(paragraph):
    """Injects a native, clickable Microsoft Word Table of Contents field."""
    run = paragraph.add_run()
    
    # 1. Begin Field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    # 2. Instruction text for TOC (Levels 1-3, Hyperlinks enabled)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)

    # 3. Separate Field
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    # 4. Placeholder text
    run2 = paragraph.add_run("Right-click here and select 'Update Field' to generate the Table of Contents.")
    run2.font.name = 'Poppins'
    run2.font.size = Pt(11)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(128, 128, 128) # Gray to indicate it needs updating

    # 5. End Field
    run3 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar3)

def create_wud_word(touchpoint_data: dict) -> BytesIO:
    """Generates a native Word doc using pure python-docx with live AWS Bedrock AI."""
    
    # 1. NULL-PROOF VARIABLE EXTRACTION
    td = touchpoint_data.get('techDetails') or {}
    wud_id = touchpoint_data.get('id') or 'TBD'
    
    raw_integration = touchpoint_data.get('integration') or 'API'
    integration_type = str(raw_integration).strip().lower()
    
    api_name = touchpoint_data.get('name') or td.get('apiName') or 'Unnamed Integration'
    target_sys = touchpoint_data.get('target') or 'Target System'
    source_sys = touchpoint_data.get('source') or 'Source System'
    owner = touchpoint_data.get('owner') or 'SDGNext Team'
    business_purpose = touchpoint_data.get('business_purpose') or 'No functional requirements listed.'
    module_name = touchpoint_data.get('module') or 'Unassigned Module'
    crm_location = touchpoint_data.get('crm_location') or 'the designated CRM process'
    
    business_flow_text = touchpoint_data.get('business_flow') or business_purpose

    # 2. ROUTE INPUTS BASED ON TYPE
    if integration_type == 'database':
        input_request_text = str(touchpoint_data.get('input') or 'Not provided')
        output_response_text = str(touchpoint_data.get('output') or 'Not provided')
    else:
        input_request_text = str(td.get('apiReq') or 'Not provided')
        output_response_text = str(td.get('apiRes') or 'Not provided')

    # ==========================================
    # LIVE AWS BEDROCK AI CALL
    # ==========================================
    ai_content = generate_wud_content(
        api_name=str(api_name), 
        module_name=str(module_name), 
        crm_location=str(crm_location), 
        business_flow=str(business_flow_text), 
        input_req=input_request_text, 
        output_res=output_response_text,
        integration_type=integration_type
    )
    
    doc = Document()

    # ==========================================
    # GLOBAL TYPOGRAPHY SETTINGS
    # ==========================================
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Poppins'
    style_normal.font.size = Pt(12)
    if style_normal.font.element.rPr is not None and style_normal.font.element.rPr.rFonts is not None:
        style_normal.font.element.rPr.rFonts.set(qn('w:asciiTheme'), '')
        style_normal.font.element.rPr.rFonts.set(qn('w:hAnsiTheme'), '')

    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Poppins'
    style_h1.font.size = Pt(16)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0, 0, 0)
    if style_h1.font.element.rPr is not None and style_h1.font.element.rPr.rFonts is not None:
        style_h1.font.element.rPr.rFonts.set(qn('w:asciiTheme'), '')
        style_h1.font.element.rPr.rFonts.set(qn('w:hAnsiTheme'), '')

    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Poppins'
    style_h2.font.size = Pt(12)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(0, 0, 0)
    if style_h2.font.element.rPr is not None and style_h2.font.element.rPr.rFonts is not None:
        style_h2.font.element.rPr.rFonts.set(qn('w:asciiTheme'), '')
        style_h2.font.element.rPr.rFonts.set(qn('w:hAnsiTheme'), '')
        
    for i in range(1, 4):
        try:
            toc_style = doc.styles[f'TOC {i}']
            toc_style.font.name = 'Poppins'
            toc_style.font.size = Pt(11)
            toc_style.font.bold = True
        except KeyError:
            pass

    # ==========================================
    # PAGE 1: TITLE PAGE
    # ==========================================
    cover_para = doc.add_paragraph()
    cover_img_path = os.path.join(os.path.dirname(__file__), 'static', 'BG_IMG.jpg')
    
    if os.path.exists(cover_img_path):
        picture = cover_para.add_run().add_picture(cover_img_path, width=Inches(8.5), height=Inches(11.7))
        set_floating_background(picture) 
    else:
        cover_para.add_run("[ Cover Image Placeholder - Please save 'BG_IMG.jpg' in app/static/ ]")

    doc.add_paragraph("\n") 
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_api = title_para.add_run(api_name)
    run_api.font.name = 'Poppins'
    run_api.font.size = Pt(26)
    run_api.font.bold = True
    run_api.font.color.rgb = RGBColor(255, 255, 255)

    type_para = doc.add_paragraph()
    type_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_type = type_para.add_run("Work Unit Document")
    run_type.font.name = 'Poppins'
    run_type.font.size = Pt(16)
    run_type.font.bold = True
    run_type.font.color.rgb = RGBColor(255, 255, 255)

    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_version = version_para.add_run("Version- 1.0")
    run_version.font.name = 'Poppins'
    run_version.font.size = Pt(14)
    run_version.font.color.rgb = RGBColor(255, 255, 255)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    current_date = datetime.now().strftime("%d-%m-%Y")
    run_date = date_para.add_run(current_date)
    run_date.font.name = 'Poppins'
    run_date.font.size = Pt(14)
    run_date.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_page_break()

    # ==========================================
    # PAGE 2: PROPRIETARY INFORMATION
    # ==========================================
    stamp_para = doc.add_paragraph()
    stamp_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    stamp_img_path = os.path.join(os.path.dirname(__file__), 'static', 'Stamp.png')
    
    if os.path.exists(stamp_img_path):
        stamp_para.add_run().add_picture(stamp_img_path, width=Inches(1.8))
    else:
        stamp_para.add_run("[ Stamp Image Placeholder - Please save 'Stamp.png' in app/static/ ]")

    copy_para = doc.add_paragraph()
    copy_run = copy_para.add_run('© Copyright 2026, Acidaes Solutions Pvt. Ltd.')
    copy_run.font.name = 'Poppins'
    copy_run.font.size = Pt(12)
    copy_run.font.bold = True
    copy_run.font.color.rgb = RGBColor(232, 31, 118)

    rights_para = doc.add_paragraph()
    rights_run = rights_para.add_run('All rights reserved.')
    rights_run.font.name = 'Poppins'
    rights_run.font.size = Pt(11)

    p1 = doc.add_paragraph()
    p1_run = p1.add_run("Contents of this document are confidential, contains ideas, concepts, processes, and other information that the Company considers proprietary. No part must be reproduced or published in any form or through any means whether electronic, mechanical, photocopying or with the aid of any information storage or retrieval system.")
    p1_run.font.name = 'Poppins'
    p1_run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2_run = p2.add_run("The material must not be disclosed to third parties without the express and prior written authorization of Acidaes Solutions Pvt. Ltd. (ASPL).")
    p2_run.font.name = 'Poppins'
    p2_run.font.size = Pt(11)
    
    p3 = doc.add_paragraph()
    p3_run = p3.add_run("Readers are to treat the information contained herein is confidential and may not disseminate copy or reproduce it in any form without the expressed written permission of the Company. 'Acidaes', ‘CRMNEXT’, ‘CRMNEXT DLP’, ‘CUSTOMERNEXT’, ‘CRMNEXT’, ‘DATANEXT’, ’ORIGINATIONNEXT’, ’BRANCHNEXT’, ’BOTNEXT’, ’Shape-Shifting Architecture’, ‘Autonoma Integration’, ‘Autobot Upgrade’, ‘Sonic Distribution System’, ‘Strategy Consonants’, ‘Pulse’ and 'Practice Led' are trademarks of Acidaes Solutions. All other trademarks are the property of their respective owners.")
    p3_run.font.name = 'Poppins'
    p3_run.font.size = Pt(11)

    doc.add_page_break()

    # ==========================================
    # PAGE 3: VERSION CONTROL
    # ==========================================
    vc_title = doc.add_paragraph()
    vc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vc_title_run = vc_title.add_run('Version Control')
    vc_title_run.font.name = 'Poppins'
    vc_title_run.font.size = Pt(14)
    
    doc.add_paragraph()

    vc_table = doc.add_table(rows=2, cols=5)
    vc_table.style = 'Table Grid'
    vc_headers = ["Sr. No.", "Date", "Version", "Description", "Author"]
    
    for i, header in enumerate(vc_headers):
        cell = vc_table.rows[0].cells[i]
        set_cell_background(cell, "D9D9D9") 
        cell.text = "" 
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.font.name = 'Poppins'
        run.font.size = Pt(11)
        run.font.bold = True
        
    vc_data = ["1", datetime.now().strftime("%d-%m-%Y"), "1.0", f"{api_name}", owner]
    for i, text in enumerate(vc_data):
        cell = vc_table.rows[1].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.name = 'Poppins'
        run.font.size = Pt(11)

    doc.add_page_break()

    # ==========================================
    # PAGE 4: TABLE OF CONTENTS
    # ==========================================
    toc_heading = doc.add_paragraph()
    toc_heading_run = toc_heading.add_run('Table of Contents')
    toc_heading_run.font.name = 'Poppins'
    toc_heading_run.font.size = Pt(16)
    toc_heading_run.font.color.rgb = RGBColor(65, 105, 225)

    toc_para = doc.add_paragraph()
    insert_dynamic_toc(toc_para)

    doc.add_page_break()

    # ==========================================
    # PAGE 5+: MAIN CONTENT
    # ==========================================
    add_custom_heading(doc, '1. Introduction', level=1)
    
    intro_para = doc.add_paragraph()
    intro_run = intro_para.add_run(ai_content.get('introduction', 'Introduction could not be generated.'))
    intro_run.font.name = 'Poppins'
    intro_run.font.size = Pt(12)

    doc.add_paragraph("\n")

    # --- Abbreviations ---
    add_custom_heading(doc, '2. Abbreviations', level=1)
    doc.add_paragraph() 
    
    master_abbreviations = {
        "API": "Application Programming Interface",
        "CRM": "Customer Relationship Management",
        "WUD": "Work Unit Document",
        "EDS": "External Data Source",
        "JSON": "JavaScript Object Notation",
        "REST": "Representational State Transfer",
        "CBS": "Core Banking System",
        "SFTP": "Secure File Transfer Protocol",
        "DB": "Database",
        "IDR": "Integration Discovery Request",
        "UI": "User Interface"
    }
    
    text_dump = f"{api_name} {source_sys} {target_sys} {json.dumps(ai_content)} {json.dumps(td)}".upper()
    found_abbs = [("1", "WUD", master_abbreviations["WUD"])] 
    sr_no = 2
    for abbr, desc in master_abbreviations.items():
        if abbr != "WUD" and abbr in text_dump:
            found_abbs.append((str(sr_no), abbr, desc))
            sr_no += 1

    abbr_table = doc.add_table(rows=1, cols=3)
    abbr_table.style = 'Table Grid'
    abbr_table.autofit = False
    abbr_table.alignment = WD_TABLE_ALIGNMENT.LEFT 
    
    abbr_table.columns[0].width = Inches(0.6)
    abbr_table.columns[1].width = Inches(1.5)
    abbr_table.columns[2].width = Inches(4.4)
    
    hdr_cells = abbr_table.rows[0].cells
    hdr_cells[0].width = Inches(0.6)
    hdr_cells[1].width = Inches(1.5)
    hdr_cells[2].width = Inches(4.4)
    
    headers = ["Sr. No.", "Abbreviation", "Description"]
    for i, header_text in enumerate(headers):
        set_cell_background(hdr_cells[i], "D9D9D9")
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(header_text)
        run.font.name = 'Poppins'
        run.font.size = Pt(11)
        run.font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    for row_data in found_abbs:
        row = abbr_table.add_row()
        row.cells[0].width = Inches(0.6)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(4.4)
        for i, val in enumerate(row_data):
            row.cells[i].text = ""
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.name = 'Poppins'
            run.font.size = Pt(11)
            if i < 2:
                row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("\n")

    # --- Technical Details ---
    add_custom_heading(doc, f'3. {api_name}', level=1)
    doc.add_paragraph(f"Source System: - {source_sys}", style='List Bullet')
    doc.add_paragraph(f"Target System: - {target_sys}", style='List Bullet')

    doc.add_paragraph("\n")

    add_custom_heading(doc, '3.1. Technical Details', level=2)
    
    tech_details_para = doc.add_paragraph()
    tech_details_run = tech_details_para.add_run("Following are the technical details:")
    tech_details_run.font.name = 'Poppins'
    tech_details_run.font.size = Pt(10)
    
    # DYNAMIC API VS DB LOGIC
    if integration_type == 'database':
        table_data = [
            ("WUD ID", str(wud_id)),
            ("Source System", str(source_sys)),
            ("Target System", str(target_sys)),
            ("Owner", str(owner)),
            ("Periodicity", "On Demand"),
            ("Interface", "Database"),
            ("Database Engine", str(td.get('dbEngine') or 'N/A')),
            ("Target Object (SP/Table)", str(td.get('dbTarget') or 'N/A')),
            ("Service Account", str(td.get('dbAccount') or 'N/A')),
            ("Firewall / IPs", str(td.get('dbFirewall') or 'N/A')),
            ("Input Data", input_request_text),
            ("Output Data", output_response_text),
            ("Failure Response", "If database exception occurs it will log and display respective error messages."),
            ("Expected Output", str(ai_content.get('expected_output') or 'NA')), 
            ("Macro Logic", str(ai_content.get('macro_logic') or 'NA')),
            ("Testing Strategy", "Database connections and SP logic will be tested as part of functional scenario testing."),
            ("Watchouts", "For unhandled exceptions, contact application administrator. Database must be accessible from target environments.")
        ]
    else:
        table_data = [
            ("WUD ID", str(wud_id)),
            ("Source System", source_sys),
            ("Target System", target_sys),
            ("Owner", owner),
            ("Authentication", td.get('apiAuth', 'N/A')),
            ("Periodicity", "On Demand"),
            ("Interface", "API"),
            ("Methodology", td.get('apiMethod', 'NA')),
            ("API Type", td.get('apiType', 'RESTFUL')),
            ("EDS Name", api_name),
            ("Input Request Payload", td.get('apiReq', 'NA')),
            ("Output Response Payload", td.get('apiRes', 'NA')),
            ("Failure Response", "If Invalid Response is passed it will display respective error messages in response."),
            ("Expected Output", ai_content.get('expected_output', 'NA')), 
            ("Macro Logic", ai_content.get('macro_logic', 'NA')),
            ("Testing Strategy", "EDS will be tested as a part of functional scenario testing.\nExplicit testing can be performed through the Widget."),
            ("Watchouts", "For unhandled exceptions, application user needs to contact application administrator.\nService must be accessible from target environments.")
        ]

    tech_table = doc.add_table(rows=1, cols=2)
    tech_table.style = 'Table Grid'
    tech_table.autofit = False
    
    for i, (key, val) in enumerate(table_data):
        if i == 0:
            row = tech_table.rows[0] 
        else:
            row = tech_table.add_row()
            
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.5)
            
        row.cells[0].text = "" 
        run_left = row.cells[0].paragraphs[0].add_run(key)
        run_left.font.name = 'Poppins'
        run_left.font.size = Pt(11)
        run_left.font.bold = True
        
        row.cells[1].text = "" 

        if key in ["Input Request Payload", "Output Response Payload", "Input Data", "Output Data"]:
            try:
                parsed_json = json.loads(val) if isinstance(val, str) else val
                formatted_val = json.dumps(parsed_json, indent=4)
            except Exception:
                formatted_val = str(val) if val else "NA"
                
            run_right = row.cells[1].paragraphs[0].add_run(formatted_val)
            run_right.font.name = 'Consolas' 
            run_right.font.size = Pt(9)      
            set_cell_background(row.cells[1], "F4F4F4") 
        else:
            val_str = str(val) if val else "NA"
            run_right = row.cells[1].paragraphs[0].add_run(val_str)
            run_right.font.name = 'Poppins'
            run_right.font.size = Pt(11)

    doc.add_paragraph("\n")
    
    # --- Conclusion ---
    add_custom_heading(doc, '4. Conclusion', level=1)
    
    conclusion_para = doc.add_paragraph()
    conclusion_run = conclusion_para.add_run(f"This WUD document formalizes the technical execution plan for the {api_name} integration.")
    conclusion_run.font.name = 'Poppins'
    conclusion_run.font.size = Pt(12)

    # ==========================================
    # FOOTERS
    # ==========================================
    for section in doc.sections:
        footer = section.footer
        footer_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
        footer_table.autofit = False
        
        footer_table.columns[0].width = Inches(1.5) 
        footer_table.columns[1].width = Inches(3.5) 
        footer_table.columns[2].width = Inches(1.5) 

        left_cell = footer_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run = left_para.add_run("Page ")
        left_run.font.name = 'Poppins'
        left_run.font.size = Pt(10)
        
        pg_run = left_para.add_run()
        pg_run.font.name = 'Poppins'
        pg_run.font.size = Pt(10)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        pg_run._r.append(fldChar1)
        pg_run._r.append(instrText)
        pg_run._r.append(fldChar2)

        mid_cell = footer_table.cell(0, 1)
        mid_para = mid_cell.paragraphs[0]
        mid_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mid_run = mid_para.add_run("www.businessnext.com")
        mid_run.font.name = 'Poppins'
        mid_run.font.size = Pt(10)
        mid_run.font.color.rgb = RGBColor(128, 128, 128)

        right_cell = footer_table.cell(0, 2)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_logo_path = os.path.join(os.path.dirname(__file__), 'static', 'logo.png')
        
        if os.path.exists(footer_logo_path):
            right_para.add_run().add_picture(footer_logo_path, height=Inches(0.35))

    settings = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)
    
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    
    return word_buffer